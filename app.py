import json
import os
import random
import io
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
import re
import traceback
import zipfile
from docx import Document
from fpdf import FPDF
from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
from sentence_transformers import SentenceTransformer, util

def clean_text_for_pdf(text):
    """
    Sanitizes string for FPDF (Standard Latin-1 encoding).
    Replaces common Unicode characters with safe equivalents and strips everything else.
    """
    if not text:
        return ""
    
    # Ensure it's a string
    text = str(text)
    
    # Common replacements for characters that crash standard FPDF fonts
    replacements = {
        '\u2013': '-', # en dash
        '\u2014': '-', # em dash
        '\u2012': '-', # figure dash
        '\u2015': '-', # horizontal bar
        '\u2212': '-', # minus sign
        '\u2018': "'", # left single quote
        '\u2019': "'", # right single quote
        '\u201a': "'", # single low-9 quotation mark
        '\u201b': "'", # single high-reversed-9 quotation mark
        '\u201c': '"', # left double quote
        '\u201d': '"', # right double quote
        '\u201e': '"', # double low-9 quotation mark
        '\u201f': '"', # double high-reversed-9 quotation mark
        '\u2022': '*', # bullet point
        '\u2026': '...', # ellipsis
        '\u20a8': 'Rs.', # Rupee symbol (old)
        '\u20b9': 'Rs.', # Rupee symbol (new)
        '\u20ac': 'EUR ', # Euro
        '\u20bf': 'B',   # Bitcoin
        '\xae': '(R)',   # Registered
        '\xa9': '(C)',   # Copyright
        '\u2122': 'TM',   # Trademark
        '\xa0': ' ',     # Non-breaking space
    }
    
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
        
    # Remove any other non-latin-1 characters to prevent FPDF crash
    safe_text = ""
    for char in text:
        try:
            char.encode('latin-1')
            safe_text += char
        except UnicodeEncodeError:
            safe_text += "?" # Fallback for unknown characters
            
    # Final safety check: ensure the string is strictly latin-1
    return safe_text.encode('latin-1', 'replace').decode('latin-1')

# ── OCR imports (graceful fallback if Tesseract binary is absent) ──────────────
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("[OCR] PyMuPDF not installed – image extraction disabled.")

try:
    from PIL import Image
    import pytesseract
    # Probe common Windows install paths for the Tesseract binary
    _TESS_PATHS = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for _p in _TESS_PATHS:
        if os.path.isfile(_p):
            pytesseract.pytesseract.tesseract_cmd = _p
            break
    # Quick smoke-test so we know it actually works
    pytesseract.get_tesseract_version()
    OCR_AVAILABLE = True
    print("[OCR] Tesseract OCR is available – image-based resume extraction enabled.")
except Exception:
    OCR_AVAILABLE = False
    print("[OCR] Tesseract not found. Install from: https://github.com/UB-Mannheim/tesseract/wiki")
# ──────────────────────────────────────────────────────────────────────────────

print("Loading AI Model (all-MiniLM-L6-v2) - This may take a minute on the first run...")
semantic_model = SentenceTransformer('all-MiniLM-L6-v2')
print("AI Model loaded successfully.")

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "fallback_secret_key_for_dev")

@app.context_processor
def inject_logos():
    return dict(COMPANY_LOGOS={
        "tata consultancy services": "tcs.jpg",
        "tcs": "tcs.jpg",
        "infosys": "infosys.png",
        "wipro": "wipro.png",
        "accenture": "accenture.png",
        "hcltech": "hcltech.png",
        "zoho corporation": "zoho.png",
        "zoho": "zoho.png",
        "flipkart": "flipkart.jpg",
        "paytm": "paytm.jpg",
        "zomato": "zomato.png",
        "amazon": "amazon.png",
        "amazon (aws)": "amazon.png",
        "aws": "amazon.png",
        "techflow": "techflow.png",
        "greengrid": "greengrid.png",
        "nexus systems": "nexussystems.png",
        "eduspark": "eduspark.jpg",
        "velociti": "velocity.jpg",
        "jpmorgan chase": "jpmorgan.png",
        "jp morgan": "jpmorgan.png",
        "figma": "figma.png",
        "canva": "canva.jpg",
        "google": "google.png",
        "google cloud": "google.png",
        "googlecloud": "google.png",
        "microsoft": "microsoft.png",
        "microsoft (azure)": "microsoft.png",
        "azure": "azure.jpg",
        "goldman sachs": "golmansachs.png",
        "morgan stanley": "Morgan_Stanley.png",
        "razorpay": "razorpay.png",
        "adobe": "adobe.png",
        "airbnb": "airbnb.png",
        "nvidia": "nvidia.png",
        "swiggy": "swiggy.jpg",
        "atlassian": "Atlassian.png",
        "deloitte": "Deloitte.png",
        "mckinsey": "McKinsey.png",
        "bcg": "bcg.png",
        "bain": "bain.png",
        "cisco": "cisco.png",
        "ibm": "ibm.png",
        "unilever": "uniliver.png",
        "p&g": "p&g.jpg",
        "fractal analytics": "Fractal_Analytics.png",
        "mu sigma": "musigma.jpg",
        "ey": "ey.png",
        "pwc": "pwc.png",
        "oracle": "oracle.jpg",
        "meta": "meta.png",
    })

# Base Jobs Database
COMPANY_DOMAINS = {
    "tata consultancy services": "tcs.com",
    "infosys": "infosys.com",
    "wipro": "wipro.com",
    "accenture": "accenture.com",
    "hcltech": "hcltech.com",
    "zoho corporation": "zoho.com",
    "flipkart": "flipkart.com",
    "paytm": "paytm.com",
    "zomato": "zomato.com",
    "amazon": "amazon.com",
    "techflow": "techflow.io",
    "greengrid": "greengrid.net",
    "nexus systems": "nexus.com",
    "eduspark": "eduspark.edu",
    "velociti": "velociti.co"
}

# Dynamic Roadmap Data Loading
ROADMAP_DOMAINS = {}
ROADMAP_COMPANIES = {}
COMPANY_ROADMAPS = {}

def load_career_roadmap():
    global ROADMAP_DOMAINS, ROADMAP_COMPANIES, COMPANY_ROADMAPS
    try:
        # 1. Load Domains (Master List)
        if os.path.exists("domains.json"):
            with open("domains.json", "r", encoding="utf-8") as f:
                domains_list = json.load(f).get("domains", [])
                
                # Helper for domain icons
                domain_icons = {
                    "Technology": "computer",
                    "Finance": "account_balance",
                    "Design": "palette",
                    "Data Science": "analytics",
                    "Cybersecurity": "shield",
                    "AI/ML": "psychology",
                    "Blockchain": "link",
                    "Product Management": "dashboard",
                    "Marketing": "campaign",
                    "Consulting": "groups",
                    "Cloud Computing": "cloud",
                    "Healthcare": "medical_services",
                    "Operations": "settings",
                    "Sales": "trending_up",
                    "Human Resources": "badge",
                    "DevOps": "terminal"
                }
                
                new_domains = {}
                for dom_name in domains_list:
                    dom_id = dom_name.lower().replace("/", "_").replace(" ", "_")
                    new_domains[dom_id] = {
                        "name": dom_name,
                        "icon": domain_icons.get(dom_name, "explore"),
                        "desc": f"Expert-curated pathways for {dom_name}."
                    }
                ROADMAP_DOMAINS = new_domains

        # 2. Load Domain Mapping & Companies
        if os.path.exists("domains_data.json"):
            with open("domains_data.json", "r", encoding="utf-8") as f:
                data = json.load(f)
                
                new_roadmap_companies = {}
                new_company_roadmaps = {}
                
                for dom_obj in data.get("domain_data", []):
                    dom_name = dom_obj["domain"]
                    dom_id = dom_name.lower().replace("/", "_").replace(" ", "_")
                    
                    if dom_id not in new_roadmap_companies:
                        new_roadmap_companies[dom_id] = []
                    
                    roles_list = dom_obj.get("roles", [])
                    
                    for comp_name in dom_obj.get("companies", []):
                        comp_id = comp_name.lower().replace(" ", "_").replace(".", "")
                        
                        new_roadmap_companies[dom_id].append({
                            "id": comp_id,
                            "name": comp_name,
                            "domain": f"{comp_name.lower().replace(' ', '')}.com"
                        })
                        
                        # Consolidate roles and skills
                        all_roles = []
                        all_skills = set()
                        for role in roles_list:
                            all_roles.append({
                                "title": role["title"],
                                "salary": role.get("salary_range_lpa", "Competitive") + " LPA"
                            })
                            for s in role.get("skills", []):
                                all_skills.add(s)
                        
                        new_company_roadmaps[comp_id] = {
                            "demands": f"{comp_name} focuses on excellence in {dom_name}. They specifically look for expertise in {', '.join(list(all_skills)[:3])}.",
                            "skills": list(all_skills),
                            "roles": all_roles,
                            "learning_resources": [] # Fallback
                        }
                
                ROADMAP_COMPANIES = new_roadmap_companies
                COMPANY_ROADMAPS = new_company_roadmaps
        else:
            print("Warning: domains_data.json not found.")
    except Exception as e:
        print(f"Error loading roadmap data: {e}")
        traceback.print_exc()

load_career_roadmap()


COMPANY_WALLPAPERS = {
    "infosys": "https://images.unsplash.com/photo-1518770660439-4636190af475?auto=format&fit=crop&w=1200&q=80",
    "tata consultancy services": "https://images.unsplash.com/photo-1451187580459-43490279c0fa?auto=format&fit=crop&w=1200&q=80",
    "amazon": "https://images.unsplash.com/photo-1523474253046-8cd2748b5fd2?auto=format&fit=crop&w=1200&q=80",
    "paytm": "https://images.unsplash.com/photo-1616077168079-7e0908d07656?auto=format&fit=crop&w=1200&q=80",
    "zomato": "https://images.unsplash.com/photo-1555396273-367ea4eb4db5?auto=format&fit=crop&w=1200&q=80",
    "flipkart": "https://images.unsplash.com/photo-1556742049-0cfed4f6a45d?auto=format&fit=crop&w=1200&q=80",
    "techflow": "https://images.unsplash.com/photo-1550751827-4bd374c3f58b?auto=format&fit=crop&w=1200&q=80",
    "greengrid": "https://images.unsplash.com/photo-1466611653911-95081537e5b7?auto=format&fit=crop&w=1200&q=80",
}

DEFAULT_WALLPAPER = "https://images.unsplash.com/photo-1504384308090-c894fdcc538d?auto=format&fit=crop&w=1200&q=80"

# --- SALARY NORMALIZATION HELPERS ---
HIGH_PAYING_THRESHOLD = 25 # Will be updated dynamically

def convert_to_inr_lpa(salary_str):
    """
    Normalizes salary strings to INR Lakhs Per Annum (LPA).
    Handles ranges, currencies ($, £, €), and timeframes (/mo, /hr).
    """
    if not salary_str or not isinstance(salary_str, str) or salary_str.lower() == "competitive":
        return "Competitive", 0
    
    salary_str = salary_str.lower().strip()
    
    # Check if already a range of numbers (LPA assumption)
    if re.match(r'^\d+(\.\d+)?\s*-\s*\d+(\.\d+)?$', salary_str):
        parts = re.findall(r'\d+(?:\.\d+)?', salary_str)
        max_val = float(parts[-1])
        return f"₹{salary_str} LPA", max_val

    # Currency Multipliers (Approximate)
    multipliers = {
        '$': 83.0,
        '£': 105.0,
        '€': 90.0,
        '₹': 1.0
    }
    
    multiplier = 1.0
    for symbol, m in multipliers.items():
        if symbol in salary_str:
            multiplier = m
            break
            
    # Extract numbers
    nums = re.findall(r'(\d+(?:\.\d+)?)\s*k?', salary_str)
    if not nums:
        return salary_str, 0
    
    max_num = float(nums[-1])
    if 'k' in salary_str:
        max_num *= 1000
        
    # Timeframe handling
    if '/ mo' in salary_str or 'monthly' in salary_str:
        max_num *= 12
    elif '/ hr' in salary_str or 'hourly' in salary_str:
        max_num *= 1920 # ~40hrs/week * 48 weeks
        
    # Convert to Lakhs Per Annum (LPA)
    lpa_val = (max_num * multiplier) / 100000
    
    if '-' in salary_str and len(nums) >= 2:
        min_num = float(nums[0])
        if 'k' in salary_str: min_num *= 1000
        if '/ mo' in salary_str or 'monthly' in salary_str: min_num *= 12
        elif '/ hr' in salary_str or 'hourly' in salary_str: min_num *= 1920
        
        lpa_min = (min_num * multiplier) / 100000
        return f"₹{lpa_min:.1f} - {lpa_val:.1f} LPA", lpa_val
    else:
        return f"₹{lpa_val:.1f} LPA", lpa_val

jobs_db = {
    "1": {"title": "Frontend Intern", "company": "TechFlow", "domain": COMPANY_DOMAINS["techflow"], "image": COMPANY_WALLPAPERS["techflow"], "skills": ["React", "Tailwind CSS", "JavaScript", "UI/UX", "Git", "TypeScript"], "location": "Palo Alto, CA", "salary": "$4k - $6k / mo", "description": "Join our core product team to build scalable React components for the next generation of SaaS tools."},
    "2": {"title": "Product Designer", "company": "GreenGrid", "domain": COMPANY_DOMAINS["greengrid"], "image": COMPANY_WALLPAPERS["greengrid"], "skills": ["Figma", "UI/UX", "Design", "Prototyping", "Accessibility"], "location": "Austin, TX", "salary": "$120k - $150k", "description": "Lead the design of complex grid management dashboards. Focus on data visualization and accessibility."},
    "3": {"title": "Data Analyst", "company": "Nexus Systems", "domain": COMPANY_DOMAINS.get("nexus systems"), "image": DEFAULT_WALLPAPER, "skills": ["SQL", "Python", "Tableau", "Excel", "Data Modeling"], "location": "London, UK", "salary": "£55k - £70k", "description": "Utilize SQL and Python to extract insights from massive datasets. Directly reporting to the CTO."},
    "4": {"title": "UX Researcher", "company": "EduSpark", "domain": COMPANY_DOMAINS.get("eduspark"), "image": DEFAULT_WALLPAPER, "skills": ["User Research", "Testing", "Interviewing", "Empathy"], "location": "Remote", "salary": "$80 - $110 / hr", "description": "Conduct user interviews and usability testing for our K-12 learning platform. Drive data-informed design."},
    "5": {"title": "Backend Engineer", "company": "Velociti", "domain": COMPANY_DOMAINS.get("velociti"), "image": DEFAULT_WALLPAPER, "skills": ["Python", "Go", "Docker", "Kubernetes", "API", "Microservices"], "location": "Berlin, DE", "salary": "€75k - €95k", "description": "Optimize high-throughput logistics algorithms. Experience with distributed systems and Kubernetes required."}
}

# Normalize original jobs_db salaries
for jid in jobs_db:
    disp, val = convert_to_inr_lpa(jobs_db[jid]["salary"])
    jobs_db[jid]["salary"] = disp
    jobs_db[jid]["salary_num"] = val

# Load extra companies and domains from new JSON files
try:
    # 1. Load Domains Data (Comprehensive)
    if os.path.exists("domains_data.json"):
        with open("domains_data.json", "r", encoding="utf-8") as f:
            d_data = json.load(f)
            idx = 100 # New ID range for imported data
            for dom_obj in d_data.get("domain_data", []):
                domain_name = dom_obj["domain"]
                for role in dom_obj.get("roles", []):
                    # For each role in each domain, add dummy entries for companies in that domain
                    for comp in dom_obj.get("companies", []):
                        comp_name_lower = comp.lower()
                        disp_sal, val_sal = convert_to_inr_lpa(role.get("salary_range_lpa", "Competitive"))
                        
                        jobs_db[str(idx)] = {
                            "title": role["title"],
                            "company": comp,
                            "domain": f"{comp.lower().replace(' ', '')}.com",
                            "image": COMPANY_WALLPAPERS.get(comp_name_lower, DEFAULT_WALLPAPER),
                            "skills": role.get("skills", []),
                            "location": "India",
                            "salary": disp_sal,
                            "salary_num": val_sal,
                            "description": f"Exciting opportunity for a {role['title']} at {comp} in the {domain_name} domain."
                        }
                        idx += 1

    # 2. Add specific companies from companies.json
    if os.path.exists("companies.json"):
        with open("companies.json", "r", encoding="utf-8") as f:
            c_data = json.load(f)
            for comp_obj in c_data.get("companies_to_apply", []):
                comp_name = comp_obj["name"]
                comp_name_lower = comp_name.lower()
                disp_sal, val_sal = convert_to_inr_lpa(comp_obj.get("average_salary_lpa", "Competitive"))
                
                # Check if we already have this company/role combo, if not add it
                for role_title in comp_obj.get("roles", []):
                    exists = any(j["company"] == comp_name and j["title"] == role_title for j in jobs_db.values())
                    if not exists:
                        jobs_db[str(idx)] = {
                            "title": role_title,
                            "company": comp_name,
                            "domain": f"{comp_name_lower.replace(' ', '')}.com",
                            "image": COMPANY_WALLPAPERS.get(comp_name_lower, DEFAULT_WALLPAPER),
                            "skills": ["General Industry Standards"],
                            "location": "Remote / India",
                            "salary": disp_sal,
                            "salary_num": val_sal,
                            "description": f"Join {comp_name} as a {role_title} and drive innovation in your field."
                        }
                        idx += 1

    # 3. Calculate Dynamic High-Paying Threshold
    salaries = [j["salary_num"] for j in jobs_db.values() if j["salary_num"] > 0]
    if salaries:
        salaries.sort()
        # Set threshold at top 30%
        threshold_idx = int(len(salaries) * 0.7)
        HIGH_PAYING_THRESHOLD = salaries[threshold_idx]
        print(f"Dynamic High-Paying Threshold set to: {HIGH_PAYING_THRESHOLD} LPA")

except Exception as e:
    print(f"Could not load extended company data: {e}")
    traceback.print_exc()

# Dummy Learning Platforms
LEARNING_PLATFORMS = {
    "Python": {"name": "Coursera: Python for Everybody", "url": "https://www.coursera.org/specializations/python"},
    "Java": {"name": "Udemy: Java Programming Masterclass", "url": "https://www.udemy.com/course/java-the-complete-java-developer-course/"},
    "React": {"name": "Frontend Masters: Complete React", "url": "https://frontendmasters.com/courses/complete-react-v8/"},
    "Docker": {"name": "Pluralsight: Docker Fundamentals", "url": "https://www.pluralsight.com/courses/docker-fundamentals"},
    "SQL": {"name": "DataCamp: Intro to SQL", "url": "https://www.datacamp.com/courses/introduction-to-sql"},
    "JavaScript": {"name": "Codecademy: JavaScript", "url": "https://www.codecademy.com/learn/introduction-to-javascript"},
    "Figma": {"name": "Udemy: Complete Web Design in Figma", "url": "https://www.udemy.com/course/complete-web-designer-mobile-web-design-in-figma/"}
}

# ── Text Cleaning Utilities ───────────────────────────────────────────────────
import unicodedata

def clean_extracted_text(raw: str) -> str:
    """
    Normalises garbled text that AI resume-builder PDFs commonly produce:
      • Unicode NFKD decomposition (ligatures → plain ASCII)
      • Collapse runs of single-spaced letters  ("P y t h o n" → "Python")
      • Normalise whitespace
    """
    # 1. Unicode normalise (handles ﬁ→fi, fancy quotes, etc.)
    text = unicodedata.normalize("NFKD", raw)

    # 2. Collapse single-char-spaced runs that template PDFs produce.
    #    Pattern: a letter, then (space + single letter) repeated 2+ times.
    #    e.g. "P y t h o n" → "Python"
    def _collapse(m):
        return m.group(0).replace(" ", "")
    text = re.sub(r'\b([A-Za-z] )(?:[A-Za-z] ){1,}[A-Za-z]\b', _collapse, text)

    # 3. Normalise whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ── Smart PDF Extraction (fitz primary, PyPDF2 fallback) ─────────────────────

def extract_text_from_pdf_smart(file_bytes: bytes) -> str:
    """
    Three-layer PDF text extraction:
      Layer 1  – PyMuPDF (fitz) native text  (best for template PDFs)
      Layer 2  – PyPDF2 fallback
      Layer 3  – OCR via PyMuPDF page-render + Tesseract
    Returns cleaned, lowercased text.
    """
    text = ""

    # --- Layer 1: fitz (handles AI template PDFs far better than PyPDF2) -----
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                text += page.get_text("text") + " "
            doc.close()
            print(f"[Extract] fitz recovered {len(text.strip())} chars.")
        except Exception as e:
            print(f"[Extract] fitz error: {e}")

    # --- Layer 2: PyPDF2 fallback (if fitz gave very little) -----------------
    if len(text.strip()) < 120:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            pdf2_text = ""
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    pdf2_text += extracted + " "
            if len(pdf2_text.strip()) > len(text.strip()):
                text = pdf2_text
                print(f"[Extract] PyPDF2 recovered {len(text.strip())} chars.")
        except Exception as e:
            print(f"[Extract] PyPDF2 error: {e}")

    # --- Layer 3: OCR as last resort -----------------------------------------
    if len(text.strip()) < 120:
        print(f"[OCR] Text still sparse ({len(text.strip())} chars). Attempting OCR...")
        if PYMUPDF_AVAILABLE and OCR_AVAILABLE:
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                ocr_text = ""
                for page in doc:
                    mat = fitz.Matrix(2, 2)  # 144 dpi
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text += pytesseract.image_to_string(img, lang="eng", config="--psm 6") + " "
                doc.close()
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    print(f"[OCR] Recovered {len(text.strip())} chars from page images.")
            except Exception as e:
                print(f"[OCR] PDF OCR error: {e}")
        elif not OCR_AVAILABLE:
            print("[OCR] Tesseract not available – cannot OCR this document.")

    return clean_extracted_text(text)


def extract_text_from_docx_smart(file_bytes: bytes) -> str:
    """
    Two-layer DOCX text extraction:
      Layer 1  – python-docx paragraph + table text
      Layer 2  – OCR on embedded images
    """
    text = ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        # Paragraphs
        text = " ".join([p.text for p in doc.paragraphs])
        # Also grab table cells (many templates put skills in tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += " " + cell.text
    except Exception as e:
        print(f"[Extract] python-docx error: {e}")

    # OCR fallback for image-heavy DOCX
    if len(text.strip()) < 120 and OCR_AVAILABLE:
        print(f"[OCR] DOCX text sparse ({len(text.strip())} chars). Running OCR on images...")
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                for name in zf.namelist():
                    if name.startswith("word/media/") and any(
                        name.lower().endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff")
                    ):
                        img = Image.open(io.BytesIO(zf.read(name)))
                        text += " " + pytesseract.image_to_string(img, lang="eng", config="--psm 6")
        except Exception as e:
            print(f"[OCR] DOCX image OCR error: {e}")

    return clean_extracted_text(text)


# --- AI PROCESSING FUNCTIONS ---

def match_skills_intelligently(resume_text, required_skills):
    """
    Hybrid skill matching:
      Pass 1 – Cleaned keyword/regex search (strict, no hallucination)
      Pass 2 – Semantic embedding backup for remaining misses (high threshold)
    """
    matched = []
    still_missing = []
    text_lower = resume_text.lower()

    # ── Pass 1: keyword matching ─────────────────────────────────────────────
    for skill in required_skills:
        skill_lower = skill.lower()
        # Try word-boundary regex first
        pattern = r'(?:^|\b|[^a-zA-Z0-9])' + re.escape(skill_lower) + r'(?:\b|[^a-zA-Z0-9]|$)'
        if re.search(pattern, text_lower):
            matched.append(skill)
        elif skill_lower in text_lower:  # plain substring fallback
            matched.append(skill)
        else:
            still_missing.append(skill)

    # ── Pass 2: semantic backup for genuinely missing skills ─────────────────
    #    Uses a HIGH threshold (0.55) so it only catches real synonyms
    #    (e.g. "JS" ↔ "JavaScript") and not project-inferred skills.
    if still_missing:
        try:
            chunks = [c.strip() for c in resume_text.split('\n') if len(c.strip()) > 10]
            if not chunks:
                chunks = [resume_text]
            resume_embeddings = semantic_model.encode(chunks, convert_to_tensor=True)
            skill_embeddings = semantic_model.encode(still_missing, convert_to_tensor=True)
            final_missing = []
            for i, skill in enumerate(still_missing):
                cos_scores = util.cos_sim(skill_embeddings[i], resume_embeddings)[0]
                best = float(cos_scores.max())
                if best >= 0.55:
                    matched.append(skill)
                    print(f"[AI] Semantic match: '{skill}' (score {best:.2f})")
                else:
                    final_missing.append(skill)
            still_missing = final_missing
        except Exception as e:
            print(f"[AI] Semantic backup error: {e}")

    return matched, still_missing

def calculate_resume_completeness(text):
    """
    Uses basic presence detection + heuristic to score completeness (30% weight).
    """
    # Target sections that are good to have
    sections = {
        "experience": ["experience", "work history", "employment", "professional background"],
        "education": ["education", "academic", "university", "college", "degree"],
        "projects": ["projects", "portfolio", "personal work", "contributions"],
        "skills": ["skills", "technologies", "tools", "expertise"],
    }
    
    found_sections = []
    missing_sections = []
    
    text_lower = text.lower()
    for section, keywords in sections.items():
        if any(kw in text_lower for kw in keywords):
            found_sections.append(section)
        else:
            missing_sections.append(section)
            
    # Calculate a ratio (0.0 to 1.0)
    score = len(found_sections) / len(sections) if sections else 0.0
    return {"found": found_sections, "missing": missing_sections, "score": score}

def get_extracted_section(text, section_name):
    """
    Attempts to extract text for a specific section (Skills, Experience, etc.) 
    by looking for common header keywords.
    """
    headers = {
        "experience": ["experience", "work history", "employment", "professional background"],
        "education": ["education", "academic", "university", "college", "degree"],
        "projects": ["projects", "portfolio", "personal work", "contributions"],
        "skills": ["skills", "technologies", "tools", "expertise"],
    }
    keywords = headers.get(section_name, [])
    text_lower = text.lower()
    
    # Find start of the section
    start_idx = -1
    found_kw = ""
    for kw in keywords:
        pos = text_lower.find(kw)
        if pos != -1:
            # Simple check: make sure it's likely a header (start of line or after newline)
            if pos == 0 or text_lower[pos-1] in ['\n', '\r', '\t', ' ']:
                start_idx = pos
                found_kw = kw
                break
    
    if start_idx == -1: 
        return "Not found."
    
    # Find the next section start to bound this one
    all_kws = [k for v in headers.values() for k in v]
    end_idx = len(text)
    for kw in all_kws:
        pos = text_lower.find(kw, start_idx + len(found_kw))
        if pos != -1 and pos < end_idx:
            # Check if it's likely a header
            if text_lower[pos-1] in ['\n', '\r', '\t', ' ']:
                end_idx = pos
            
    extracted = text[start_idx:end_idx].strip()
    # Remove the header itself if it's just the keyword
    if extracted.lower().startswith(found_kw):
        extracted = extracted[len(found_kw):].strip(": \n\r")
        
    return extracted.capitalize()

def generate_feedback(matched, missing, completeness, final_score):
    """
    Generates intelligent structured feedback components.
    """
    feedback = {
        "summary": "",
        "strengths": [],
        "improvements": []
    }
    
    if final_score < 50:
        feedback["summary"] = f"Alignment Score: {final_score}%. The resume does not strongly align with this role's requirements."
    elif final_score < 75:
        feedback["summary"] = f"Alignment Score: {final_score}%. The resume shows potential, but is missing core requirements."
    else:
        feedback["summary"] = f"Alignment Score: {final_score}%. Outstanding! The resume is highly aligned with this position."
        
    if matched:
        feedback["strengths"].append(f"Successfully highlighted these required skills: {', '.join(matched)}.")
    
    if missing:
        feedback["improvements"].append(f"Missing core skills to work on: {', '.join(missing)}.")
    else:
        feedback["strengths"].append("Impressive mapping of all demanded technical skills.")
        
    if "projects" in completeness["missing"]:
        feedback["improvements"].append("Add a 'Projects' section to showcase practical applications of your skills.")
    if "experience" in completeness["missing"]:
        feedback["improvements"].append("Flesh out your 'Experience' timeline to demonstrate professional impact.")
        
    return feedback

def generate_dynamic_roadmap(missing_skills, completeness):
    """
    Generates personalized next-steps based on gaps found during analysis.
    """
    roadmap = []
    
    # Structural steps
    if "projects" in completeness["missing"]:
        roadmap.append("Build 2-3 real-world projects showcasing your abilities.")
    if "education" in completeness["missing"] or "experience" in completeness["missing"]:
        roadmap.append("Flesh out your resume with clear Education & Experience timelines.")

    # Skill steps
    for skill in missing_skills[:3]: # Limit to top 3 to avoid overwhelming
        roadmap.append(f"Learn the fundamentals of {skill} and add it to your skill stack.")
        
    if not roadmap:
        roadmap.append("Keep refining your current expertise to land advanced roles!")
        
    return roadmap

# ── COMPANY PRIORITIZATION ──
POPULAR_COMPANIES = [
    "Google", "Microsoft", "Amazon", "Adobe", "Meta", "Nvidia", 
    "Goldman Sachs", "JP Morgan", "McKinsey", "BCG", "Bain",
    "Flipkart", "Swiggy", "Zomato", "Atlassian", "Razorpay"
]

def get_prioritized_jobs():
    """
    Returns jobs_db items sorted by company popularity/relevance.
    """
    sorted_items = sorted(
        jobs_db.items(),
        key=lambda item: (
            0 if item[1]["company"] in POPULAR_COMPANIES else 1, # Popular first
            -item[1].get("salary_num", 0) # Then by salary
        )
    )
    return dict(sorted_items)

@app.route("/")
def dashboard():
    # Load domains for navigation
    domains = []
    try:
        if os.path.exists("domains.json"):
            with open("domains.json", "r", encoding="utf-8") as f:
                domains = json.load(f).get("domains", [])
    except Exception as e:
        print(f"Error loading domains.json: {e}")

    all_prioritized = get_prioritized_jobs()
    # Only show top 6 for home page
    home_jobs = dict(list(all_prioritized.items())[:6])
    
    return render_template("dashboard.html", jobs=home_jobs, domains=domains)

@app.route("/companies")
def companies():
    all_prioritized = get_prioritized_jobs()
    return render_template("companies.html", jobs=all_prioritized)

@app.route("/job/<job_id>")
def job_detail(job_id):
    job = jobs_db.get(job_id)
    if not job:
        return "Job not found", 404
    return render_template("job_detail.html", job_id=job_id, job=job)

@app.route("/upload/<job_id>", methods=["POST"])
def upload_resume(job_id):
    if "resume" not in request.files:
        flash("Please upload a file.")
        return redirect(url_for("job_detail", job_id=job_id))
    
    file = request.files["resume"]
    if file.filename == "":
        flash("No file selected.")
        return redirect(url_for("job_detail", job_id=job_id))
        
    allowed_exts = (".pdf", ".docx", ".doc")
    if not file.filename.lower().endswith(allowed_exts):
        flash("Error: Only PDF and DOCX resume files are allowed.")
        return redirect(url_for("job_detail", job_id=job_id))
    
    # ── Smart Extraction ──────────────────────────────────────────────────────
    filename_lower = file.filename.lower()
    file_bytes = file.read()  # Read once into memory so we can reuse

    if filename_lower.endswith(".pdf"):
        text = extract_text_from_pdf_smart(file_bytes)
    elif filename_lower.endswith((".docx", ".doc")):
        text = extract_text_from_docx_smart(file_bytes)
    else:
        text = ""

    text = text.lower()
    print(f"[DEBUG] Extracted text preview (first 500 chars): {text[:500]}")
    # ─────────────────────────────────────────────────────────────────────────
    
    # Basic Validation Check
    completeness = calculate_resume_completeness(text)
    if completeness['score'] == 0.0 and len(text) > 10:
        flash("AI validation issue: Document does not contain typical resume sections. Please upload an actual resume.")
        return redirect(url_for("job_detail", job_id=job_id))
    
    # 1. Semantic Skill Matching -> Intelligent Word Matching
    job = jobs_db.get(job_id)
    job_skills_original = job.get("skills", [])
    
    if not job_skills_original:
        job_skills_original = ["Communication", "Teamwork", "Problem Solving"]
        
    matched, missing = match_skills_intelligently(text, job_skills_original)
    
    # 2. Weighted Scoring Calculation
    skill_match_percentage = (len(matched) / len(job_skills_original)) if job_skills_original else 0.0
    
    # Weight formula: 70% skills, 30% completeness
    weighted_score_raw = (skill_match_percentage * 0.70) + (completeness['score'] * 0.30)
    match_percentage = int(weighted_score_raw * 100)
    
    # 3. Intelligent Feedback & Roadmap Generation
    feedback = generate_feedback(matched, missing, completeness, match_percentage)
    dynamic_roadmap = generate_dynamic_roadmap(missing, completeness)
    
    # Determine outcome
    selected = match_percentage >= 50 # Adjusted slightly for weighted model

    # Storing results in session along with new AI-generated fields
    session[f"analysis_{job_id}"] = {
        "match_percentage": match_percentage,
        "matched_skills": matched,
        "missing_skills": missing,
        "selected": selected,
        "feedback": feedback,
        "dynamic_roadmap": dynamic_roadmap,
        "completeness": completeness,
        "resume_text": text
    }

    return redirect(url_for("analysis", job_id=job_id))

@app.route("/analysis/<job_id>")
def analysis(job_id):
    job = jobs_db.get(job_id)
    return render_template("analysis.html", job_id=job_id, job=job)

@app.route("/api/analyze/<job_id>")
def run_analysis(job_id):
    analysis_data = session.get(f"analysis_{job_id}")
    # Fallback to random if directly accessed without uploading
    if not analysis_data:
        selected = random.choice([True, False])
        if selected:
            return {"status": "complete", "redirect": url_for("selected", job_id=job_id)}
        else:
            return {"status": "complete", "redirect": url_for("not_selected", job_id=job_id)}

    if analysis_data["selected"]:
        return {"status": "complete", "redirect": url_for("selected", job_id=job_id)}
    else:
        return {"status": "complete", "redirect": url_for("not_selected", job_id=job_id)}

@app.route("/selected/<job_id>")
def selected(job_id):
    job = jobs_db.get(job_id)
    analysis_data = session.get(f"analysis_{job_id}", {})
    
    # Map platforms for any missing skills (even selected candidates may have gaps)
    learning_suggs = []
    for missing in analysis_data.get("missing_skills", []):
        default_sugg = {"name": f"Udemy: Complete {missing} Bootcamp", "url": f"https://www.udemy.com/courses/search/?src=ukw&q={missing}"}
        sugg = LEARNING_PLATFORMS.get(missing, default_sugg)
        learning_suggs.append({"skill": missing, "platform": sugg["name"], "url": sugg["url"]})
    
    return render_template("selected.html", job_id=job_id, job=job, analysis=analysis_data, platforms=learning_suggs)

@app.route("/not-selected/<job_id>")
def not_selected(job_id):
    job = jobs_db.get(job_id)
    analysis_data = session.get(f"analysis_{job_id}", {})
    
    # Map platforms
    learning_suggs = []
    for missing in analysis_data.get("missing_skills", []):
        default_sugg = {"name": f"Udemy: Complete {missing} Bootcamp", "url": f"https://www.udemy.com/courses/search/?src=ukw&q={missing}"}
        sugg = LEARNING_PLATFORMS.get(missing, default_sugg)
        learning_suggs.append({"skill": missing, "platform": sugg["name"], "url": sugg["url"]})
    
    return render_template("not_selected.html", job_id=job_id, job=job, analysis=analysis_data, platforms=learning_suggs)

@app.route("/roadmap/<job_id>")
def roadmap(job_id):
    job = jobs_db.get(job_id)
    analysis_data = session.get(f"analysis_{job_id}", {})
    
    learning_suggs = []
    for missing in analysis_data.get("missing_skills", []):
        default_sugg = {"name": f"Udemy: Complete {missing} Bootcamp", "url": f"https://www.udemy.com/courses/search/?src=ukw&q={missing}"}
        sugg = LEARNING_PLATFORMS.get(missing, default_sugg)
        learning_suggs.append({"skill": missing, "platform": sugg["name"], "url": sugg["url"]})

    # Prepare domain data for the interactive JS section
    domain_data = {}
    for dom_id, dom_info in ROADMAP_DOMAINS.items():
        comp_names = [c["name"] for c in ROADMAP_COMPANIES.get(dom_id, [])]
        
        # Collect unique skills across all companies in this domain for the preview
        domain_skills = set()
        for comp_obj in ROADMAP_COMPANIES.get(dom_id, []):
            comp_id = comp_obj["id"]
            for skill in COMPANY_ROADMAPS.get(comp_id, {}).get("skills", []):
                domain_skills.add(skill)
        
        domain_data[dom_id] = {
            "title": dom_info["name"],
            "desc": dom_info["desc"],
            "companies": comp_names[:4], # limit for UI
            "skills": list(domain_skills)[:4] # limit for UI
        }

    return render_template("roadmap.html", job_id=job_id, job=job, analysis=analysis_data, platforms=learning_suggs, domain_data=domain_data)


@app.route("/download_report/<job_id>")
def download_report(job_id):
    analysis_data = session.get(f"analysis_{job_id}")
    job = jobs_db.get(job_id)
    
    if not analysis_data or not job:
        return "Report data missing. Please upload your resume first.", 404

    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # ── Color Palette ───────────────────────────────────────────────────
        PRIMARY_BLUE = (0, 74, 198)
        BG_COLOR = (240, 244, 248)
        TEXT_GRAY = (100, 116, 139)
        SUCCESS_GREEN = (5, 150, 105)
        ERROR_RED = (220, 38, 38)
        # ────────────────────────────────────────────────────────────────────

        # BG Top Header Decoration
        pdf.set_fill_color(*PRIMARY_BLUE)
        pdf.rect(0, 0, 210, 45, 'F')

        # LOGO 1: Company Logo (Top Left)
        comp_name_lower = job['company'].lower()
        company_logo_map = inject_logos()["COMPANY_LOGOS"]
        logo_filename = company_logo_map.get(comp_name_lower)
        if logo_filename:
            logo_path = os.path.join("static", "images", logo_filename)
            if os.path.exists(logo_path):
                # Draw logo in a white box effect
                pdf.set_fill_color(255, 255, 255)
                pdf.rect(10, 10, 25, 25, 'F')
                pdf.image(logo_path, 11, 11, 23, 23)

        # LOGO 2: Nexus Brand (Top Right - Favicon + Text)
        favicon_path = os.path.join("static", "images", "favicon.png")
        if os.path.exists(favicon_path):
            pdf.image(favicon_path, 165, 10, 10, 10)
            pdf.set_xy(176, 10)
            pdf.set_font("Arial", 'B', 10)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(20, 10, txt=clean_text_for_pdf("NEXUS"), ln=False)
            pdf.set_xy(176, 14)
            pdf.set_font("Arial", '', 7)
            pdf.cell(20, 10, txt=clean_text_for_pdf("TALENT AI"), ln=False)

        # Title Content (Over Blue Header)
        pdf.set_xy(0, 12)
        pdf.set_font("Arial", 'B', 22)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(0, 10, txt=clean_text_for_pdf("Analysis Report"), ln=True, align='C')
        
        pdf.set_font("Arial", '', 11)
        pdf.cell(0, 8, txt=clean_text_for_pdf(f"{job['title']} | {job['company']}"), ln=True, align='C')
        
        pdf.set_y(48) # Move below blue header

        # ── Score Visualization ──────────────────────────────────────────────
        match_pct = analysis_data.get('match_percentage', 0)
        
        pdf.set_font("Arial", 'B', 14)
        pdf.set_text_color(*PRIMARY_BLUE)
        pdf.cell(0, 10, txt=clean_text_for_pdf("Overall Alignment Score"), ln=True)
        
        # Draw Progress Bar
        pdf.set_fill_color(230, 230, 230) # Background gray
        pdf.rect(10, pdf.get_y(), 190, 8, 'F')
        
        # Success color for progress
        if match_pct >= 70: pdf.set_fill_color(*SUCCESS_GREEN)
        elif match_pct >= 40: pdf.set_fill_color(245, 158, 11) # Orange
        else: pdf.set_fill_color(*ERROR_RED)
        
        bar_width = (match_pct / 100) * 190
        pdf.rect(10, pdf.get_y(), bar_width, 8, 'F')
        
        pdf.set_xy(10, pdf.get_y() + 2)
        pdf.set_font("Arial", 'B', 10)
        pdf.set_text_color(255, 255, 255)
        if bar_width > 20: # only draw text inside if space
            pdf.cell(bar_width, 4, txt=f"{match_pct}%", align='R')
        pdf.ln(10)
        
        # Status Box
        status = "Recommended" if analysis_data.get('selected') else "Limited Match"
        pdf.set_font("Arial", 'B', 12)
        pdf.set_text_color(255, 255, 255)
        color = SUCCESS_GREEN if analysis_data.get('selected') else ERROR_RED
        pdf.set_fill_color(*color)
        pdf.cell(50, 8, txt=clean_text_for_pdf(status), ln=True, align='C', fill=True)
        pdf.ln(5)

        # ── Factors that built the score ──
        completeness = analysis_data.get('completeness', {})
        found_sections = completeness.get('found', [])
        matched_skills = analysis_data.get('matched_skills', [])
        
        pdf.set_font("Arial", 'B', 10)
        pdf.set_text_color(*TEXT_GRAY)
        pdf.cell(0, 8, txt=clean_text_for_pdf("What built this score:"), ln=True)
        pdf.set_font("Arial", '', 9)
        pdf.set_text_color(80, 80, 80)
        factors = []
        if matched_skills: factors.append(f"{len(matched_skills)} Relevant Skills")
        if found_sections: factors.append(f"{len(found_sections)} Valid Resume Sections")
        pdf.cell(0, 5, txt=clean_text_for_pdf(" + ".join(factors)), ln=True)
        pdf.ln(5)
        # ────────────────────────────────────────────────────────────────────

        # Structured Feedback Sections
        feedback = analysis_data.get('feedback', {})
        summary = feedback.get("summary", "") if isinstance(feedback, dict) else str(feedback)
        strengths = feedback.get("strengths", []) if isinstance(feedback, dict) else []
        improvements = feedback.get("improvements", []) if isinstance(feedback, dict) else []

        # Executive Summary "Card"
        pdf.set_font("Arial", 'B', 14)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 10, txt=clean_text_for_pdf("Executive Summary"), ln=True)
        pdf.set_font("Arial", '', 11)
        pdf.set_text_color(50, 50, 50)
        pdf.multi_cell(pdf.epw, 7, txt=clean_text_for_pdf(summary))
        pdf.ln(5)

        # Strengths Section
        pdf.set_font("Arial", 'B', 13)
        pdf.set_text_color(*SUCCESS_GREEN)
        pdf.cell(0, 10, txt=clean_text_for_pdf("Key Strengths & Matched Skills"), ln=True)
        pdf.set_font("Arial", '', 11)
        pdf.set_text_color(0, 0, 0)
        
        if analysis_data.get('matched_skills'):
            pdf.set_font("Arial", 'B', 10)
            pdf.cell(0, 6, txt=clean_text_for_pdf(f"Matched Skills: {', '.join(analysis_data['matched_skills'])}"), ln=True)
            pdf.set_font("Arial", '', 11)
            
        for s in strengths:
            pdf.multi_cell(pdf.epw, 7, txt=clean_text_for_pdf(f"• {s}"))
        pdf.ln(5)

        # Improvements Section
        pdf.set_font("Arial", 'B', 13)
        pdf.set_text_color(*ERROR_RED)
        pdf.cell(0, 10, txt=clean_text_for_pdf("Gaps & Improvements Needed"), ln=True)
        pdf.set_font("Arial", '', 11)
        pdf.set_text_color(0, 0, 0)
        
        if analysis_data.get('missing_skills'):
            pdf.set_font("Arial", 'B', 10)
            pdf.cell(0, 6, txt=clean_text_for_pdf(f"Skills Gaps: {', '.join(analysis_data['missing_skills'])}"), ln=True)
            pdf.set_font("Arial", '', 11)

        for imp in improvements:
            pdf.multi_cell(pdf.epw, 7, txt=clean_text_for_pdf(f"• {imp}"))
        
        # ── APPENDIX: EXTRACTED RESUME DATA ──────────────────────────────────
        pdf.add_page()
        pdf.set_fill_color(*BG_COLOR)
        pdf.rect(0, 0, 210, 297, 'F') # Light background for appendix
        
        pdf.set_y(15)
        pdf.set_font("Arial", 'B', 18)
        pdf.set_text_color(*PRIMARY_BLUE)
        pdf.cell(0, 12, txt=clean_text_for_pdf("Resume Content Summary"), ln=True)
        pdf.set_font("Arial", 'I', 10)
        pdf.set_text_color(*TEXT_GRAY)
        pdf.cell(0, 6, txt=clean_text_for_pdf("Raw data extracted from uploaded document for verification"), ln=True)
        pdf.ln(10)
        
        resume_text = analysis_data.get("resume_text", "")
        
        # Extract and print sections
        for section in ["skills", "experience", "projects"]:
            pdf.set_font("Arial", 'B', 14)
            pdf.set_text_color(*PRIMARY_BLUE)
            pdf.cell(0, 10, txt=clean_text_for_pdf(section.upper()), ln=True)
            
            pdf.set_font("Arial", '', 10)
            pdf.set_text_color(20, 20, 20)
            content = get_extracted_section(resume_text, section)
            pdf.multi_cell(pdf.epw, 6, txt=clean_text_for_pdf(content))
            pdf.ln(10)
            
        # ────────────────────────────────────────────────────────────────────

        pdf.set_y(-25)
        pdf.set_font("Arial", 'I', 8)
        pdf.set_text_color(150, 150, 150)
        pdf.cell(0, 10, txt=clean_text_for_pdf("Generated Autonomously by Nexus AI Core. Confidential candidate profile."), ln=True, align='C')
        
        pdf_bytes = pdf.output()
        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"Detailed_Analysis_{comp_name_lower}.pdf"
        )
    except Exception as e:
        print(f"[ERROR] PDF Generation failed: {str(e)}")
        traceback.print_exc()
        return f"Error generating premium report: {str(e)}. Please contact support.", 500


def generate_resume_pdf(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, txt=clean_text_for_pdf(f"{data['name']}"), ln=True, align='C')
    
    pdf.set_font("Arial", '', 12)
    contact_info = f"Email: {data.get('email', 'N/A')} | Phone: {data.get('phone', 'N/A')}"
    pdf.cell(0, 10, txt=clean_text_for_pdf(contact_info), ln=True, align='C')
    pdf.cell(0, 10, txt=clean_text_for_pdf(f"Domain: {data.get('domain', 'N/A')}"), ln=True, align='C')
    pdf.ln(10)
    
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, txt=clean_text_for_pdf("Professional Summary"), ln=True)
    pdf.set_font("Arial", '', 12)
    summary = data.get('summary', 'Motivated professional with a strong desire to learn and grow in the industry.')
    pdf.multi_cell(pdf.epw, 10, txt=clean_text_for_pdf(summary))
    pdf.ln(5)
    
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, txt=clean_text_for_pdf("Key Skills"), ln=True)
    pdf.set_font("Arial", '', 12)
    skills = data.get('skills', 'Problem Solving, Communication, Teamwork')
    pdf.multi_cell(pdf.epw, 10, txt=clean_text_for_pdf(skills))
    pdf.ln(5)

    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, txt=clean_text_for_pdf("Experience / Education Highlight"), ln=True)
    pdf.set_font("Arial", '', 12)
    experience = data.get('experience', 'Relevant background experience making me a strong candidate for future roles.')
    pdf.multi_cell(pdf.epw, 10, txt=clean_text_for_pdf(experience))

    return pdf.output()

def generate_resume_docx(data):
    doc = Document()
    doc.add_heading(data['name'], 0)
    p = doc.add_paragraph()
    p.add_run(f"Email: {data.get('email', 'N/A')} | Phone: {data.get('phone', 'N/A')}").bold = True
    doc.add_paragraph(f"Domain: {data.get('domain', 'N/A')}")
    
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(data.get('summary', 'Motivated professional with a strong desire to learn and grow in the industry.'))
    
    doc.add_heading('Key Skills', level=1)
    doc.add_paragraph(data.get('skills', 'Problem Solving, Communication, Teamwork'))
    
    doc.add_heading('Experience / Education Highlight', level=1)
    doc.add_paragraph(data.get('experience', 'Relevant background experience making me a strong candidate for future roles.'))
    
    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)
    return mem


@app.route("/roadmap_domains")
def roadmap_domains():
    return render_template("roadmap_domains.html", domains=ROADMAP_DOMAINS)

@app.route("/roadmap_domain/<domain_id>")
def roadmap_domain(domain_id):
    domain = ROADMAP_DOMAINS.get(domain_id)
    if not domain:
        return "Domain not found", 404
    companies = ROADMAP_COMPANIES.get(domain_id, [])
    return render_template("roadmap_companies.html", domain_id=domain_id, domain=domain, companies=companies)

@app.route("/company_roadmap_detail/<company_id>")
def company_roadmap_detail(company_id):
    company_info = None
    for dom, comps in ROADMAP_COMPANIES.items():
        for comp in comps:
            if comp["id"] == company_id:
                company_info = comp
                break
        if company_info:
            break
    if not company_info:
        return "Company not found", 404
    
    roadmap_data = COMPANY_ROADMAPS.get(company_id, {
        "demands": "Details for this company are currently being updated.",
        "skills": ["General Industry Standards"],
        "roles": [{"title": "Various Roles", "salary": "Competitive"}],
        "learning_resources": []
    })
    
    return render_template("company_roadmap_detail.html", company=company_info, roadmap=roadmap_data)

@app.route("/resume_builder", methods=["GET", "POST"])
def resume_builder():
    if request.method == "POST":
        data = {
            "name": request.form.get("name"),
            "email": request.form.get("email"),
            "phone": request.form.get("phone"),
            "domain": request.form.get("domain"),
            "summary": request.form.get("summary"),
            "skills": request.form.get("skills"),
            "experience": request.form.get("experience"),
        }
        session["generated_resume_data"] = data
        return redirect(url_for("resume_ready"))
    return render_template("resume_builder.html")

@app.route("/resume_ready")
def resume_ready():
    data = session.get("generated_resume_data")
    if not data:
        return redirect(url_for("resume_builder"))
    return render_template("resume_ready.html", data=data)

@app.route("/download_resume/<fmt>")
def download_resume(fmt):
    data = session.get("generated_resume_data")
    if not data:
        return redirect(url_for("resume_builder"))
    
    if fmt == 'pdf':
        pdf_bytes = generate_resume_pdf(data)
        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"{data['name'].replace(' ', '_')}_Resume.pdf"
        )
    elif fmt == 'docx':
        docx_io = generate_resume_docx(data)
        return send_file(
            docx_io,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"{data['name'].replace(' ', '_')}_Resume.docx"
        )
    return "Invalid format", 400

if __name__ == "__main__":
    is_debug = os.getenv("DEBUG", "True").lower() == "true"
    port = int(os.getenv("PORT", 5000))
    app.run(debug=is_debug, port=port)
